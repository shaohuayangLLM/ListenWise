import io
from datetime import datetime

from docx import Document as DocxDocument
from docx.shared import Pt, Inches
from fpdf import FPDF

from app.models.base import SceneType

SCENE_LABELS = {
    SceneType.requirement_review: "需求评审会",
    SceneType.report_meeting: "汇报会议",
    SceneType.leadership_conference: "领导会议",
    SceneType.parent_meeting: "家长会",
    SceneType.phone_call: "电话通话",
    SceneType.study_recording: "学习录音",
}


def _format_duration(seconds: int) -> str:
    h, remainder = divmod(seconds, 3600)
    m, s = divmod(remainder, 60)
    if h > 0:
        return f"{h}小时{m}分钟"
    return f"{m}分钟{s}秒"


def _format_timestamp(ts: float) -> str:
    m, s = divmod(int(ts), 60)
    return f"{m:02d}:{s:02d}"


def _build_metadata_lines(title: str, scene_type: SceneType, duration: int, created_at: datetime) -> list[str]:
    return [
        f"标题: {title}",
        f"场景: {SCENE_LABELS.get(scene_type, scene_type.value)}",
        f"时长: {_format_duration(duration)}",
        f"创建时间: {created_at.strftime('%Y-%m-%d %H:%M')}",
    ]


def _build_document_sections(content: dict) -> list[tuple[str, str]]:
    """Convert document content JSON into (heading, body) pairs."""
    sections = []

    if summary := content.get("summary"):
        sections.append(("摘要", summary))

    if key_decisions := content.get("key_decisions"):
        if isinstance(key_decisions, list):
            body = "\n".join(f"- {d}" for d in key_decisions)
        else:
            body = str(key_decisions)
        sections.append(("关键决策", body))

    if action_items := content.get("action_items"):
        if isinstance(action_items, list):
            body = "\n".join(f"- {item}" for item in action_items)
        else:
            body = str(action_items)
        sections.append(("待办事项", body))

    if discussion_points := content.get("discussion_points"):
        if isinstance(discussion_points, list):
            body = "\n".join(f"- {p}" for p in discussion_points)
        else:
            body = str(discussion_points)
        sections.append(("讨论要点", body))

    if meeting_info := content.get("meeting_info"):
        if isinstance(meeting_info, dict):
            body = "\n".join(f"- {k}: {v}" for k, v in meeting_info.items())
        else:
            body = str(meeting_info)
        sections.append(("会议信息", body))

    # Handle any remaining keys not already processed
    known_keys = {"summary", "key_decisions", "action_items", "discussion_points", "meeting_info"}
    for key, value in content.items():
        if key in known_keys:
            continue
        heading = key.replace("_", " ").title()
        if isinstance(value, list):
            body = "\n".join(f"- {item}" for item in value)
        elif isinstance(value, dict):
            body = "\n".join(f"- {k}: {v}" for k, v in value.items())
        else:
            body = str(value)
        sections.append((heading, body))

    return sections


def _build_transcript_text(segments: list | dict, include_timestamps: bool) -> str:
    """Convert transcript segments to readable text."""
    if isinstance(segments, dict):
        segments = segments.get("segments", [])
    if not isinstance(segments, list):
        return ""

    lines = []
    for seg in segments:
        if not isinstance(seg, dict):
            continue
        speaker = seg.get("speaker", "")
        text = seg.get("text", "")
        if include_timestamps:
            start = _format_timestamp(seg.get("start", 0))
            prefix = f"[{start}]"
            if speaker:
                prefix += f" {speaker}:"
            lines.append(f"{prefix} {text}")
        else:
            if speaker:
                lines.append(f"{speaker}: {text}")
            else:
                lines.append(text)

    return "\n".join(lines)


# --- Markdown Export ---

def export_markdown(
    title: str,
    scene_type: SceneType,
    duration: int,
    created_at: datetime,
    document_content: dict | None,
    transcript_segments: list | dict | None,
    transcript_full_text: str | None,
    include_transcript: bool = True,
    include_timestamps: bool = True,
) -> bytes:
    lines = [f"# {title}", ""]

    for meta_line in _build_metadata_lines(title, scene_type, duration, created_at):
        lines.append(f"> {meta_line}")
    lines.append("")

    # Document sections
    if document_content:
        lines.append("---")
        lines.append("")
        for heading, body in _build_document_sections(document_content):
            lines.append(f"## {heading}")
            lines.append("")
            lines.append(body)
            lines.append("")

    # Transcript
    if include_transcript:
        lines.append("---")
        lines.append("")
        lines.append("## 转录文本")
        lines.append("")
        if transcript_segments:
            lines.append(_build_transcript_text(transcript_segments, include_timestamps))
        elif transcript_full_text:
            lines.append(transcript_full_text)
        else:
            lines.append("*暂无转录文本*")
        lines.append("")

    return "\n".join(lines).encode("utf-8")


# --- Word (DOCX) Export ---

def export_docx(
    title: str,
    scene_type: SceneType,
    duration: int,
    created_at: datetime,
    document_content: dict | None,
    transcript_segments: list | dict | None,
    transcript_full_text: str | None,
    include_transcript: bool = True,
    include_timestamps: bool = True,
) -> bytes:
    doc = DocxDocument()

    # Title
    doc.add_heading(title, level=0)

    # Metadata
    for meta_line in _build_metadata_lines(title, scene_type, duration, created_at):
        p = doc.add_paragraph(meta_line)
        p.style = doc.styles["No Spacing"]
        for run in p.runs:
            run.font.size = Pt(10)
    doc.add_paragraph("")

    # Document sections
    if document_content:
        for heading, body in _build_document_sections(document_content):
            doc.add_heading(heading, level=1)
            for line in body.split("\n"):
                if line.startswith("- "):
                    doc.add_paragraph(line[2:], style="List Bullet")
                else:
                    doc.add_paragraph(line)

    # Transcript
    if include_transcript:
        doc.add_heading("转录文本", level=1)
        if transcript_segments:
            text = _build_transcript_text(transcript_segments, include_timestamps)
            for line in text.split("\n"):
                p = doc.add_paragraph(line)
                for run in p.runs:
                    run.font.size = Pt(10)
        elif transcript_full_text:
            doc.add_paragraph(transcript_full_text)
        else:
            doc.add_paragraph("暂无转录文本")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# --- PDF Export ---

def _find_cjk_font() -> str | None:
    """Find a CJK-capable TTF font on the system."""
    import platform
    from pathlib import Path

    candidates = []
    if platform.system() == "Darwin":
        candidates = [
            "/System/Library/Fonts/STHeiti Medium.ttc",
            "/System/Library/Fonts/STHeiti Light.ttc",
            "/System/Library/Fonts/PingFang.ttc",
            "/Library/Fonts/Arial Unicode.ttf",
        ]
    elif platform.system() == "Linux":
        candidates = [
            "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
            "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
            "/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf",
        ]

    for path in candidates:
        if Path(path).exists():
            return path
    return None


class _ChinesePDF(FPDF):
    def __init__(self):
        super().__init__()
        self.set_auto_page_break(auto=True, margin=15)
        self._font_family_name = "Helvetica"

        font_path = _find_cjk_font()
        if font_path:
            self.add_font("CJK", "", font_path, uni=True)
            self._font_family_name = "CJK"

    def _use_font(self, style: str = "", size: int = 10):
        if self._font_family_name == "CJK":
            # CJK fonts loaded via add_font don't support B/I styles
            self.set_font(self._font_family_name, "", size)
        else:
            self.set_font(self._font_family_name, style, size)

    def header(self):
        self._use_font("I", 8)
        self.cell(0, 10, "ListenWise", align="R", new_x="LMARGIN", new_y="NEXT")

    def chapter_title(self, title: str):
        self._use_font("B", 14)
        self.cell(0, 10, title, new_x="LMARGIN", new_y="NEXT")
        self.ln(2)

    def section_title(self, title: str):
        self._use_font("B", 12)
        self.cell(0, 8, title, new_x="LMARGIN", new_y="NEXT")
        self.ln(2)

    def body_text(self, text: str):
        self._use_font("", 10)
        self.multi_cell(0, 5, text)
        self.ln(2)

    def meta_text(self, text: str):
        self._use_font("I", 9)
        self.cell(0, 5, text, new_x="LMARGIN", new_y="NEXT")


def export_pdf(
    title: str,
    scene_type: SceneType,
    duration: int,
    created_at: datetime,
    document_content: dict | None,
    transcript_segments: list | dict | None,
    transcript_full_text: str | None,
    include_transcript: bool = True,
    include_timestamps: bool = True,
) -> bytes:
    pdf = _ChinesePDF()
    pdf.add_page()

    # Title
    pdf.chapter_title(title)

    # Metadata
    for meta_line in _build_metadata_lines(title, scene_type, duration, created_at):
        pdf.meta_text(meta_line)
    pdf.ln(5)

    # Document sections
    if document_content:
        for heading, body in _build_document_sections(document_content):
            pdf.section_title(heading)
            pdf.body_text(body)

    # Transcript
    if include_transcript:
        pdf.section_title("Transcript")
        if transcript_segments:
            text = _build_transcript_text(transcript_segments, include_timestamps)
            pdf.body_text(text)
        elif transcript_full_text:
            pdf.body_text(transcript_full_text)
        else:
            pdf.body_text("No transcript available.")

    return pdf.output()
